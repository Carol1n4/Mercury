import re
import pythoncom
from docx import Document
from docx.shared import Pt, Inches, Mm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import Flask, request, redirect, url_for, render_template, send_file, flash
import os
import tempfile
from docx import Document
from pdf2docx import Converter
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from flask_sqlalchemy import SQLAlchemy
from docx2pdf import convert  # Para convertir el archivo .docx a .pdf
import pythoncom
import resumen2
from werkzeug.security import generate_password_hash, check_password_hash
import secrets
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docx.shared import Mm

app = Flask(__name__)
app.secret_key = secrets.token_hex(16)

# Configuración de la base de datos SQLite
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///usuarios.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

# Modelo de Usuario
class Usuario(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(100), unique=True, nullable=False)
    password = db.Column(db.String(100), nullable=False)
    codigo = db.Column(db.String(50), nullable=False) 
    institucion = db.Column(db.String(100), nullable=False)

# Modelo para los Códigos Institucionales Permitidos
class CodigoPermitido(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    codigo = db.Column(db.String(50), unique=True, nullable=False)  # códigos únicos permitidos

# Crear la base de datos y las tablas
with app.app_context():
    db.create_all()
    
@app.route('/', methods=['GET', 'POST'])
def registro():
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        codigo = request.form.get('codigo')
        institucion = request.form.get('institucion')

        # Comprobar si el usuario ya existe
        usuario_existente = Usuario.query.filter_by(email=email).first()
        if usuario_existente:
            flash('El correo ya está registrado', 'error')
            return redirect(url_for('registro'))
        
        # Verificar si el código institucional ya está registrado
        codigo_permitido = CodigoPermitido.query.filter_by(codigo=codigo).first()
        if not codigo_permitido:
            flash('El código institucional no es válido.', 'error')
            return redirect(url_for('registro'))  # Asegúrate de no permitir la creación de usuarios con códigos no válidos

        # Crear un nuevo usuario
        hashed_password = generate_password_hash(password)
        nuevo_usuario = Usuario(email=email, password=hashed_password, codigo=codigo, institucion=institucion)
        db.session.add(nuevo_usuario)
        db.session.commit()
        return redirect(url_for('login'))  # Redirigir al login tras registrarse

    return render_template('registro.html')


# Ruta para mostrar el formulario de inicio de sesión
@app.route('/Login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']

        # Buscar si el usuario existe en la base de datos
        usuario_existente = Usuario.query.filter_by(email=email).first()

        if not usuario_existente:
            flash('El usuario no existe.', 'error')
        
        elif not check_password_hash(usuario_existente.password, password):
            flash('Contraseña incorrecta.', 'error')
        
        else:
            flash(f'Bienvenido, {email}!', 'success')
            return redirect(url_for('upload'))  

    return render_template('login.html')



@app.route('/upload', methods=['GET', 'POST'])
def upload():
    if request.method == 'POST':
        file = request.files.get('file')
        if not file or file.filename == '':
            return 'No file part or no selected file'
        
        # Guardar el archivo temporalmente
        nombre_original, extension = os.path.splitext(file.filename)
        ruta_archivo = os.path.join(tempfile.gettempdir(), nombre_original + extension)
        file.save(ruta_archivo)

        dificultad_aprendizaje = request.form.get('dificultad_aprendizaje')
        nombre_fuente = request.form.get('nombre_fuente')
        interlineado = int(request.form.get('interlineado', 14))
        size_fuente = int(request.form.get('size_fuente', 12))
        formato_salida = request.form.get('formato_salida', 'pdf')
        color_fondo = request.form.get('color_fondo')

        if dificultad_aprendizaje == 'TDAH':
            # Llamar a la función de resumen de `resumen2.py`
            resumen_generado = resumen2.generar_resumen(ruta_archivo)
            if resumen_generado:
                return send_file(resumen_generado, as_attachment=True, download_name = f"{nombre_original}_tdah.pdf")
            else:
                return 'Error al generar el resumen.', 500
        elif dificultad_aprendizaje == 'Dislexia':
            # Adaptar el archivo para dislexia
            archivo_adaptado = modificar_word(ruta_archivo, nombre_fuente=nombre_fuente, interlineado=interlineado, 
                                                   size_fuente=size_fuente, formato_salida=formato_salida, color_fondo=color_fondo)
            if archivo_adaptado:
                return send_file(archivo_adaptado, as_attachment=True)
            else:
                return 'Error al adaptar el documento.', 500
        elif dificultad_aprendizaje == 'Ambas':
            resumen_generado = resumen2.generar_resumen(ruta_archivo)
            if resumen_generado:
                # Luego aplicar modificaciones para dislexia al resumen
                archivo_adaptado = modificar_word(resumen_generado, nombre_fuente=nombre_fuente, interlineado=interlineado, 
                                                    size_fuente=size_fuente,formato_salida=formato_salida, color_fondo=color_fondo)
                if archivo_adaptado:
                    return send_file(archivo_adaptado, as_attachment=True, download_name=f"{nombre_original}_tdah_dislexia.pdf")
                else:
                    return 'Error al adaptar el resumen para dislexia.', 500
            else:
                return 'Error al generar el resumen.', 500
    return render_template('upload.html')


def convertir_pdf_a_docx(ruta_pdf, ruta_docx):
    cv = Converter(ruta_pdf)
    cv.convert(ruta_docx, start=0, end=None)
    cv.close()

def detectar_latex(texto):
    """
    Función para detectar y separar secciones de LaTeX del resto del texto.
    
    :param texto: Cadena de texto con posibles fórmulas LaTeX.
    :return: Lista de tuplas con el tipo de segmento (latex o texto) y el contenido.
    """
    # Patrón regex para encontrar secciones de LaTeX en $...$, \[...\], \begin{...}...\end{...}
    patron_latex = re.compile(
        r'(\$.*?\$)'                       # Expresiones matemáticas en $...$
        r'|(\$\$.*?\$\$)'                 # Expresiones matemáticas en $$...$$
        r'|\\\[(.*?)\\\]'                 # Expresiones matemáticas en \[...\]
        r'|\\begin\{.*?\}(.*?)\\end\{.*?\}'  # Entornos \begin{...}...\end{...}
        , re.DOTALL
    )
    
    segmentos = []
    ultimo_indice = 0

    # Buscar coincidencias de LaTeX en el texto
    for coincidencia in patron_latex.finditer(texto):
        inicio, fin = coincidencia.span()

        # Agregar texto normal antes de la coincidencia
        if inicio > ultimo_indice:
            segmentos.append(("texto", texto[ultimo_indice:inicio]))

        # Agregar la coincidencia de LaTeX
        segmentos.append(("latex", texto[inicio:fin]))
        ultimo_indice = fin

    # Agregar cualquier texto restante
    if ultimo_indice < len(texto):
        segmentos.append(("texto", texto[ultimo_indice:]))

    return segmentos

def adaptar_texto(texto, nombre_fuente, interlineado, size_fuente):
    """
    Aplica adaptaciones específicas de dislexia y TDAH a un texto.
    """
    # Aquí podrías personalizar la lógica de adaptación de texto para dislexia o TDAH.
    # Por ejemplo, aumentar la fuente, cambiar la fuente, etc.
    # Este es un ejemplo simple de cambio de fuente y tamaño.
    return texto.upper()  # Ejemplo básico

def aplicar_adaptaciones(parrafo, nombre_fuente, interlineado, size_fuente):
    """
    Función para aplicar adaptaciones a los segmentos que no son LaTeX.
    
    :param parrafo: Objeto Paragraph en el que se aplica la adaptación.
    """
    segmentos = detectar_latex(parrafo.text)
    
    nuevo_contenido = ""
    
    for tipo, contenido in segmentos:
        if tipo == "latex":
            nuevo_contenido += contenido  # Mantener LaTeX intacto
        else:
            nuevo_contenido += adaptar_texto(contenido, nombre_fuente, interlineado, size_fuente)
    
    # Borrar el texto existente y reemplazarlo por el nuevo contenido
    parrafo.clear()
    parrafo.add_run(nuevo_contenido)

def modificar_word(ruta_archivo, nombre_fuente, interlineado, size_fuente, color_fondo):
    pythoncom.CoInitialize()  # Inicializa el sistema COM
    try:
        doc = Document(ruta_archivo)
        
        # Aplicar el color de fondo a las páginas de Word
        sectPr = doc.sections[0]._sectPr
        bg = OxmlElement('w:background')
        bg.set(qn('w:color'), color_fondo[1:])  # Eliminamos el símbolo '#' para obtener el color HEX puro
        sectPr.append(bg)

        for section in doc.sections:
            section.page_width = Mm(210)  # A4 width en milímetros
            section.page_height = Mm(297)  # A4 height en milímetros
            section.top_margin = Pt(72)  # Ajusta el margen superior
            section.bottom_margin = Inches(0.5)  # Ajusta el margen inferior
            section.left_margin = Inches(0.8)  # Ajusta el margen izquierdo
            section.right_margin = Inches(0.8)  # Ajusta el margen derecho

        for parrafo in doc.paragraphs:
            parrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            aplicar_adaptaciones(parrafo, nombre_fuente, interlineado, size_fuente)

            p = parrafo._element
            pPr = p.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:line'), str(int((interlineado + 2) * 1.8 * 20)))
            pPr.append(spacing)

        nombre_archivo_modificado = os.path.splitext(ruta_archivo)[0] + "_Adaptado.docx"
        doc.save(nombre_archivo_modificado)
        return nombre_archivo_modificado
    finally:
        pythoncom.CoUninitialize()  # Desinicializa COM
